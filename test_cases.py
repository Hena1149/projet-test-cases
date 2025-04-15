from docx import Document
import random

from nlp_processing import tokenize_and_lemmatize
def generate_test_case(rule, index):
    """Génère un cas de test structuré"""
    return [
        f"CT-{index:03d}",
        f"Vérifier que {rule}",
        "Aucune précondition spécifique",
        f"1. Exécuter {rule}\n2. Vérifier le résultat",
        f"Le système doit respecter: {rule}",
        "Non exécuté"
    ]

def save_test_cases_to_docx(test_cases, filename):
    """Sauvegarde dans un document Word"""
    doc = Document()
    doc.add_heading('Cas de Test', 0)
    
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "ID", "Description"
    hdr[2].text, hdr[3].text = "Préconditions", "Étapes"
    hdr[4].text, hdr[5].text = "Résultat", "Statut"
    
    for case in test_cases:
        row = table.add_row().cells
        for i, val in enumerate(case):
            row[i].text = str(val)
    
    doc.save(filename)